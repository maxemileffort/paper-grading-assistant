# Adapted from https://github.com/robsalgado/personal_data_science_projects/blob/master/topic_modeling_nmf/nlp_topic_utils.ipynb

import glob, os, re, shutil, string, sys
from joblib import dump, load
from PIL import Image
from sklearn import preprocessing
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS, CountVectorizer
from sklearn.svm import SVC
import nltk
from nltk.tokenize import TweetTokenizer, RegexpTokenizer
nltk.download('stopwords')
nltk.download('wordnet')

from nltk.stem.wordnet import WordNetLemmatizer
import pandas as pd
import numpy as np
import streamlit as st
import docx
import fitz  

import pytesseract
from pytesseract import Output, TesseractError


# Contraction map
c_dict = {
    "ain't": "am not",
    "aren't": "are not",
    "can't": "cannot",
    "can't've": "cannot have",
    "'cause": "because",
    "could've": "could have",
    "couldn't": "could not",
    "couldn't've": "could not have",
    "didn't": "did not",
    "doesn't": "does not",
    "don't": "do not",
    "hadn't": "had not",
    "hadn't've": "had not have",
    "hasn't": "has not",
    "haven't": "have not",
    "he'd": "he would",
    "he'd've": "he would have",
    "he'll": "he will",
    "he'll've": "he will have",
    "he's": "he is",
    "how'd": "how did",
    "how'd'y": "how do you",
    "how'll": "how will",
    "how's": "how is",
    "i'd": "I would",
    "i'd've": "I would have",
    "i'll": "I will",
    "i'll've": "I will have",
    "i'm": "I am",
    "i've": "I have",
    "isn't": "is not",
    "it'd": "it had",
    "it'd've": "it would have",
    "it'll": "it will",
    "it'll've": "it will have",
    "it's": "it is",
    "let's": "let us",
    "ma'am": "madam",
    "mayn't": "may not",
    "might've": "might have",
    "mightn't": "might not",
    "mightn't've": "might not have",
    "must've": "must have",
    "mustn't": "must not",
    "mustn't've": "must not have",
    "needn't": "need not",
    "needn't've": "need not have",
    "o'clock": "of the clock",
    "oughtn't": "ought not",
    "oughtn't've": "ought not have",
    "shan't": "shall not",
    "sha'n't": "shall not",
    "shan't've": "shall not have",
    "she'd": "she would",
    "she'd've": "she would have",
    "she'll": "she will",
    "she'll've": "she will have",
    "she's": "she is",
    "should've": "should have",
    "shouldn't": "should not",
    "shouldn't've": "should not have",
    "so've": "so have",
    "so's": "so is",
    "that'd": "that would",
    "that'd've": "that would have",
    "that's": "that is",
    "there'd": "there had",
    "there'd've": "there would have",
    "there's": "there is",
    "they'd": "they would",
    "they'd've": "they would have",
    "they'll": "they will",
    "they'll've": "they will have",
    "they're": "they are",
    "they've": "they have",
    "to've": "to have",
    "wasn't": "was not",
    "we'd": "we had",
    "we'd've": "we would have",
    "we'll": "we will",
    "we'll've": "we will have",
    "we're": "we are",
    "we've": "we have",
    "weren't": "were not",
    "what'll": "what will",
    "what'll've": "what will have",
    "what're": "what are",
    "what's": "what is",
    "what've": "what have",
    "when's": "when is",
    "when've": "when have",
    "where'd": "where did",
    "where's": "where is",
    "where've": "where have",
    "who'll": "who will",
    "who'll've": "who will have",
    "who's": "who is",
    "who've": "who have",
    "why's": "why is",
    "why've": "why have",
    "will've": "will have",
    "won't": "will not",
    "won't've": "will not have",
    "would've": "would have",
    "wouldn't": "would not",
    "wouldn't've": "would not have",
    "y'all": "you all",
    "y'alls": "you alls",
    "y'all'd": "you all would",
    "y'all'd've": "you all would have",
    "y'all're": "you all are",
    "y'all've": "you all have",
    "you'd": "you had",
    "you'd've": "you would have",
    "you'll": "you you will",
    "you'll've": "you you will have",
    "you're": "you are",
    "you've": "you have"
}

# Compiling the contraction dict
c_re = re.compile('(%s)' % '|'.join(c_dict.keys()))

# List of stop words
add_stop = ['said', 'say', '...', 'like', 'ad', 'ha', 'wa', 'reuters', 'just', 'cap']
stop_words = ENGLISH_STOP_WORDS.union(add_stop)

# List of punctuation
punc = list(set(string.punctuation))

# Splits words on white spaces (leaves contractions intact) and splits out
# trailing punctuation
def casual_tokenizer(text):
    """ Called in process_text. """
    tokenizer = TweetTokenizer()
    tokens = tokenizer.tokenize(text)
    return tokens

def expand_contractions(text, c_re=c_re):
    """ Called in process_text. """
    def replace(match):
        return c_dict[match.group(0)]
    return c_re.sub(replace, text)

def process_text(text):
    text = casual_tokenizer(text)
    text = [each.lower() for each in text]
    text = [expand_contractions(each, c_re=c_re) for each in text]
    text = [re.sub('[0-9]+', '', each) for each in text]
    text = [re.sub(r"(@\[A-Za-z0-9]+)|([^0-9A-Za-z \t])|(\w+:\/\/\S+)|^rt|http.+?", "", each) for each in text]
    text = [re.sub('[^a-zA-Z]', ' ', each) for each in text]
    text = [WordNetLemmatizer().lemmatize(each) for each in text]
    text = [w for w in text if w not in punc]
    text = [w for w in text if w not in stop_words]
    text = [each for each in text if len(each) > 1]
    text = [each for each in text if ' ' not in each]
    # text = unique_words(text)
    return text

def top_words(topic, n_top_words):
    return topic.argsort()[:-n_top_words - 1:-1]  

def topic_table(model, feature_names, n_top_words):
    topics = {}
    for topic_idx, topic in enumerate(model.components_):
        t = (topic_idx)
        topics[t] = [feature_names[i] for i in top_words(topic, n_top_words)]
    return pd.DataFrame(topics)

def whitespace_tokenizer(text): 
    pattern = r"(?u)\b\w\w+\b" 
    tokenizer_regex = RegexpTokenizer(pattern)
    tokens = tokenizer_regex.tokenize(text)
    return tokens

# Funtion to remove duplicate words
def unique_words(text): 
    ulist = []
    [ulist.append(x) for x in text if x not in ulist]
    return ulist

def word_count(text):
    return len(str(text).split(' '))

def process_df(data, max_score):
    """ Turns essays into a DF that can be used to train a classificatio model. """
    data['word_count'] = data.essay.apply(word_count)
    data['page_count'] = data.word_count.apply(lambda x: str(round(x/250, 1)).rstrip('0'))
    data['tokenized_essay'] = data.essay.apply(process_text)
    data['token_count'] = data.tokenized_essay.apply(len)
    data = data.fillna(0)
    data['max_score'] = max_score
    data['letter_grade'] = 0
    return data

def create_models(df):
    # gensim model
    try:
        from gensim.models import KeyedVectors
        gensim_model = KeyedVectors.load('./models/mywordvecs.kvmodel')
    except:
        import gensim.downloader as api
        gensim_model = api.load("glove-wiki-gigaword-300")
        gensim_model.save('./models/mywordvecs.kvmodel')
    # tf model
    # print(df)
    tokenized_text = df['tokenized_essay']
    try:
        tf_vectorizer = load('./models/kaggle_pretrained_tf_model.joblib')
    except:
        no_features = 1000
        tf_vectorizer = CountVectorizer(max_df=0.85, 
                                        min_df=1, 
                                        max_features=no_features, )
                                        #stop_words='english', 
                                        # preprocessor=' '.join)
    tf = tf_vectorizer.fit_transform(tokenized_text)
    y = df['class']
    # clf model
    try:
        clf = load('./models/kaggle_pretrained_clf_model.joblib')
    except:
        clf = SVC()
        clf.fit(tf,y)
    # save tf and clf
    dump(tf_vectorizer, './models/kaggle_pretrained_tf_model.joblib')
    dump(clf, './models/kaggle_pretrained_clf_model.joblib')

def get_gensim_model():
    from gensim.models import KeyedVectors
    gensim_model = KeyedVectors.load('./models/mywordvecs.kvmodel')
    return gensim_model

def get_clf_model():
    clf_model = load('./models/kaggle_pretrained_clf_model.joblib')
    return clf_model

def get_tf_model():
    tf_model = load('./models/kaggle_pretrained_tf_model.joblib')
    return tf_model

def generate_para_topics(df):
    """ Takes a tokenized paragraph and labels it with a topic. """
    from sklearn.feature_extraction.text import CountVectorizer
    from sklearn.decomposition import LatentDirichletAllocation

    # use or build w2v model
    df = df[df['token_p'].map(len) > 5]
    # print(df)
    tokenized_text = df['token_p']
    # LDA can only use raw term counts for LDA because it is a probabilistic graphical model
    no_features = 1000
    tf_vectorizer = CountVectorizer(max_df=0.85, 
                                    min_df=1, 
                                    max_features=no_features, 
                                    #stop_words='english', 
                                    preprocessor=' '.join)
    try:
        tf = tf_vectorizer.fit_transform(tokenized_text) # tf embeddings
    except:
        return 'Problem with file.'
    # print("tf: \n", tf)
    tf_feature_names = tf_vectorizer.get_feature_names()

    n_components = 10

    # Run LDA
    lda = LatentDirichletAllocation(n_components=n_components, max_iter=5, 
                                    learning_method='online', learning_offset=50.,
                                    random_state=0).fit(tf)

    lda_docweights = lda.transform(tf_vectorizer.transform(tokenized_text))

    n_top_words = 10

    lda_topic_df = topic_table(
        lda,
        tf_feature_names,
        n_top_words
    ).T

    # Cleaning up the top words to create topic summaries
    lda_topic_df['topics'] = lda_topic_df.apply(lambda x: [' '.join(x)], axis=1) # Joining each word into a list
    lda_topic_df['topics'] = lda_topic_df['topics'].str[0]  # Removing the list brackets
    lda_topic_df['topics'] = lda_topic_df['topics'].apply(lambda x: whitespace_tokenizer(x)) # tokenize
    lda_topic_df['topics'] = lda_topic_df['topics'].apply(lambda x: unique_words(x))  # Removing duplicate words
    lda_topic_df['topics'] = lda_topic_df['topics'].apply(lambda x: [' '.join(x)])  # Joining each word into a list
    lda_topic_df['topics'] = lda_topic_df['topics'].str[0]  # Removing the list brackets

    lda_topic_df = lda_topic_df['topics'].reset_index()
    lda_topic_df.columns = ['lda_topic_num', 'topics']

    # Creating a temp df with the id and topic num to join on
    id_ = df['para_id'].tolist()
    df_temp = pd.DataFrame({
        'para_id': id_,
        'lda_topic_num': lda_docweights.argmax(axis=1)
    })
    merged_topic = df_temp.merge(
        lda_topic_df,
        on='lda_topic_num',
        how='left'
    )
    # Merging with the original df
    df_topics = pd.merge(
        df,
        merged_topic,
        on='para_id',
        how='left'
    )

    # set top label with probabilities
    # borrowed from https://stackoverflow.com/questions/35252762/finding-number-of-documents-per-topic-for-lda-with-scikit-learn
    df_topics['top_label'] = 'n/a'
    docsVStopics = pd.DataFrame(lda_docweights, 
                                columns=["Topic "+str(i) for i in range(n_top_words)])
    most_likely_topics = docsVStopics.idxmax(axis=1)
    df_topics['top_label'] = most_likely_topics
    label_arr = []
    for i in range(len(df_topics)):
        idx = int(df_topics.iloc[i]["top_label"][-1])
        arr = (df_topics.iloc[i]["topics"]).split(" ")
        label_arr.append(arr[idx])
    df_topics['top_label'] = label_arr

    return df_topics

def empty_data_folder():
    # base_dir = '/app/paper-grading-assistant/streamlit'
    # try:
    #     if st.secrets['prod_env']:
    #         dir_name = base_dir + "/data/"
    #     else:
    #         dir_name = "./data/"
    # except:
    #         dir_name = "./data/"
    dir_name = "./data"
    files = os.listdir(dir_name)
    for item in files:
        try:
            shutil.rmtree(os.path.join(dir_name, item))
        except:
            os.remove(dir_name+'/'+item)
    return

def save_uploaded_file(uploadedfile):
    dir_name = "./data"
    filename = os.path.join(dir_name,uploadedfile.name)
    with open(filename,"wb") as f:
        f.write(uploadedfile.getbuffer())
    return filename

def handle_pdf(filename):
    with fitz.open(filename) as doc:
        full_text = ""
        zoom = 1.2
        mat = fitz.Matrix(zoom, zoom)
        noOfPages = doc.pageCount 
        for pageNo in range(noOfPages):
            page = doc.load_page(pageNo) 
            pix = page.get_pixmap(matrix = mat) 
            output = './data/' + str(pageNo) + '.jpg'
            pix.save(output) 

            text = str(pytesseract.image_to_string(Image.open(output)))
            full_text += text
            os.remove(output)
        return full_text

def get_file_text(filename):
    full_text = []
    if filename.endswith('.docx'):
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            full_text.append(para.text)
    elif filename.endswith('.pdf'):
        text = handle_pdf(filename)
        # print("pdf_text:", text)
        return text
    return '\n'.join(full_text)

def essay2df(essay):
    paragraphs = essay.split('\n')
    df = pd.DataFrame(paragraphs, columns=['paragraphs'])
    return df

def extract_papers(filename):
    # base_dir = '/app/paper-grading-assistant/streamlit'
    # try:
    #     if st.secrets['prod_env']:
    #         dir_name = base_dir + "/data/"
    #     else:
    #         dir_name = "./data/"
    # except:
    #         dir_name = "./data/"
    dir_name = "./data/"
    import zipfile
    with zipfile.ZipFile(filename, 'r') as zip_ref:
        zip_ref.extractall(dir_name)
    # delete zip and txt file
    files = os.listdir(dir_name)
    print("files:", files)
    for item in files:
        if item.endswith(".zip") or item.endswith(".txt"):
            os.remove(os.path.join(dir_name, item))
    files = os.listdir(dir_name)
    # go into path and create dataframe out of docx and pdf files
    if len(files) == 1:
        # sometimes the files are nested one more folder deep:
        path_ = os.path.join(dir_name, os.listdir(dir_name)[0])
    else: 
        # if the files extract properly:
        path_ = dir_name
    print(path_)
    files = os.listdir(path_)
    doc_lst = []
    for item in files:
        document = get_file_text(path_+'/'+item)
        doc_lst.append([item, document])
    df = pd.DataFrame(doc_lst, columns=['file', 'essay'])
    print(df)
    return df

def swap_grade(grade):
    if grade == 5:
        grade = 'A'
    elif grade == 4:
        grade = 'B'
    elif grade == 3:
        grade = 'C'
    elif grade == 2:
        grade = 'D'
    else:
        grade = 'F'
    return grade

def fix_org_score(score):
    comment = ''
    if score >= 0.5:
        comment = "Great paper with just some small tweaks."
    elif score >= 0.25 and score < 0.5:
        comment = "Clear organizational structure, but could use some work."
    elif score < 0.25 and score >= 0:
        comment = "Organization makes the paper difficult to read."
    else:
        comment = "Something went wrong here. Maybe this isn't a paper, or it was submitted in an unsupported format."
    return comment

def grade_papers(uploadedfile, max_score, new_model=False): 
    # first, empty data folder
    empty_data_folder()
    # save file 
    filename = save_uploaded_file(uploadedfile)
    # extract contents
    df = extract_papers(filename)
    df['essay_id'] = range(1,len(df)+1)
    # pre-process data
    df = process_df(df, max_score)
    # train models
    if new_model:
        # TODO replace this with the teacher's model
        clf, tf_vec = train_teachers_models()
    else:
        clf, tf_vec = get_pretrained_models()
    # predict grades
    tf = tf_vec.transform(df.tokenized_essay.apply(" ".join))
    y_pred = clf.predict(tf)
    df['letter_grade'] = y_pred
    df['org_score'] = 0
    # run through topic modeling model
    lst_coefs=[]
    for i in range(len(df)):
        p = essay2df(df.essay.iloc[i])
        p['para_id'] = range(1,len(p)+1)
        p['word_count'] = p.paragraphs.apply(word_count)
        p['token_p'] = p.paragraphs.apply(process_text)
        p['token_count'] = p.token_p.apply(len)
        p_topics = generate_para_topics(p)
        if isinstance(p_topics, str):
            lst_coefs.append(-100)
            continue
        high_num_same_top = p_topics['top_label'].value_counts()[0]
        p_topics_len = len(p_topics)
        freq = high_num_same_top / p_topics_len
        # print("frequency of topic with highest occurrences:", freq)
        model = get_gensim_model()
        lst_sims = []
        for i in range(p_topics_len):
            try:
                # print(p_topics['top_label'].iloc[i], p_topics['top_label'].iloc[i+1])
                sim = model.n_similarity(p_topics['top_label'].iloc[i], p_topics['top_label'].iloc[i+1])
                # print(sim)
                lst_sims.append(sim)
            except:
                # print(p_topics['top_label'].iloc[i], p_topics['top_label'].iloc[0])
                sim = model.n_similarity(p_topics['top_label'].iloc[i], p_topics['top_label'].iloc[0])
                # print(sim)
                lst_sims.append(sim)
        sim_mean = np.array(lst_sims).mean()
        # print("mean of similarities:", sim_mean)
        
        # relating frequency of the most common topic to the 
        # mean of similarities bewtween all of the topics

        # most common topic * mean of similarities between paragraph topics. 
        # that way, if they are both high, the score is high. if one is 
        # low, there's a middle range, and if they are both low, then that 
        # shows big problems. multiplied so theres a bigger range of values vs addition.
        my_coef = round(freq * sim_mean, 4)
        # print("my_coef:", my_coef)
        lst_coefs.append(my_coef) 
    df['org_score'] = lst_coefs
    # drop columns teachers won't use
    drop_cols = ['tokenized_essay', 'token_count', 'max_score', 'essay_id']
    df = df.drop(columns=drop_cols)
    # swap letters for numbers on letter_grade column
    df['letter_grade'] = df['letter_grade'].apply(swap_grade)
    # make sense of org_score
    df['org_score'] = df['org_score'].apply(fix_org_score)
    df = df.rename(columns={'org_score': "Organization Score", 
                            "letter_grade":"Letter Grade",
                            "file":"File",
                            "page_count":"Page Count",
                            "essay":"Essay",
                            "word_count":"Word Count"})
    # empty data folder again
    empty_data_folder()
    return df

def train_teachers_models(t_df):
    # TODO finish this
    t_df = replicate_csv(t_df)
    tf_vectorizer = CountVectorizer(max_df=0.85, 
                                min_df=2, 
                                max_features=1000, 
                                stop_words='english') 
    tf = tf_vectorizer.fit_transform(t_df['tokenized_essay'])
    X_tf = tf
    y = t_df['letter_grade']

def get_pretrained_models():
    clf = load('./models/kaggle_pretrained_clf_model.joblib')
    tf_vec = load('./models/kaggle_pretrained_tf_model.joblib')
    return clf, tf_vec

def replicate_csv(data):
    # recreating csv from modeling_ii notebook
    cols_to_drop = [col for col in data.columns if 'rater' in col]
    data = data.drop(columns=cols_to_drop)
    data['tokenized_essay'] = data.essay.apply(process_text)
    data['word_count'] = data.essay.apply(word_count)
    data = data.fillna(0)
    data['max_score'] = 0
    essay_sets = data.essay_set.unique()
    for set_ in essay_sets:
        if set_ == 1:
            data.loc[data.essay_set == set_, 'max_score'] = 12
        if set_ == 2:
            data.loc[data.essay_set == set_, 'max_score'] = 10
        if set_ == 3 or set_ == 4:
            data.loc[data.essay_set == set_, 'max_score'] = 3
        if set_ == 5 or set_ == 6:
            data.loc[data.essay_set == set_, 'max_score'] = 4
        if set_ == 7:
            data.loc[data.essay_set == set_, 'max_score'] = 30
        if set_ == 8:
            data.loc[data.essay_set == set_, 'max_score'] = 60
    data['pct_score'] = 0
    for set_ in essay_sets:
        if set_ == 2:
            data.loc[data.essay_set == set_, 'pct_score'] = (data.loc[data.essay_set==set_,'domain1_score'] \
                                                       + data.loc[data.essay_set==set_,'domain2_score']) \
                                                       / data.loc[data.essay_set==set_,'max_score']
            continue
        else:
            data.loc[data.essay_set == set_, 'pct_score'] = data.loc[data.essay_set==set_,'domain1_score'] \
                                                       / data.loc[data.essay_set==set_,'max_score']
    data['class'] = 1
    for x in range(len(data)):
        if (data.pct_score[x]) >= .9:
            data['class'][x] = 5
            continue
        elif data.pct_score[x] >= .8 and data.pct_score[x] < .9:
            data['class'][x] = 4
            continue
        elif data.pct_score[x] >= .7 and data.pct_score[x] < .8:
            data['class'][x] = 3
            continue
        elif data.pct_score[x] >= .6 and data.pct_score[x] < .7:
            data['class'][x] = 2
    return data

def set_config():
    base_dir = '/app/paper-grading-assistant/streamlit'
    try:
        im = Image.open(base_dir+"/images/984102_avatar_casual_male_man_person_icon.ico")
    except:
        im = Image.open("./images/984102_avatar_casual_male_man_person_icon.ico")
    st.set_page_config(
        page_title="I'm Skip, your grading assistant.",
        page_icon=im,
    )
    return

def setup_folders():
    from pathlib import Path

    folder_names = ["data", 
                    "models",
                    'sample_data',
                    'images'
                    ]

    for folder in folder_names:
        _file = Path(f'./{folder}')
        if _file.exists():
            pass
        else:
            os.mkdir(f'./{folder}')
    
    url='https://raw.githubusercontent.com/maxemileffort/paper-grading-assistant/master/streamlit/sample_data/processed_essays.csv'
    df = pd.read_csv(url, sep=",", error_bad_lines=False, header=0, index_col=0)
    # print(df)
    create_models(df)
    # set_config()
    st.session_state['models_loaded'] = True
    try:
        empty_data_folder()
    except:
        pass